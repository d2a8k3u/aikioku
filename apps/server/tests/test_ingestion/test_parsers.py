"""Tests for ingestion parsers (PDF, DOCX, audio, image, web, email)."""

from __future__ import annotations

import builtins
import io
from unittest.mock import MagicMock, patch

import pytest

from src.models.note import Note


# ---------------------------------------------------------------------------
# Helpers to generate real file bytes inside Docker where deps are installed
# ---------------------------------------------------------------------------


def _make_pdf_bytes(title: str = "PDF Title", body: str = "PDF body text.") -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), f"{title}\n\n{body}")
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return buf.read()


def _make_docx_bytes(title: str = "DOCX Title", body: str = "DOCX body text.") -> bytes:
    import docx

    document = docx.Document()
    document.add_paragraph(title)
    document.add_paragraph(body)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------


class TestPdfParser:
    def test_parse_pdf_success(self):
        from src.ingestion.pdf_parser import parse_pdf

        pdf_bytes = _make_pdf_bytes("Hello PDF", "This is the content.")
        note = parse_pdf(pdf_bytes, "test.pdf")
        assert isinstance(note, Note)
        assert note.title == "Hello PDF"
        assert "This is the content." in note.content
        assert note.path == "test.pdf"

    def test_import_error_when_pymupdf_missing(self):
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "fitz":
                raise ImportError("No module named 'fitz'")
            return real_import(name, *args, **kwargs)

        with patch.object(builtins, "__import__", fake_import):
            from src.ingestion.pdf_parser import parse_pdf

            with pytest.raises(ImportError):
                parse_pdf(b"fake pdf", "test.pdf")


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------


class TestDocxParser:
    def test_parse_docx_success(self):
        from src.ingestion.docx_parser import parse_docx

        docx_bytes = _make_docx_bytes("Hello DOCX", "This is the content.")
        note = parse_docx(docx_bytes, "test.docx")
        assert isinstance(note, Note)
        assert note.title == "Hello DOCX"
        assert "This is the content." in note.content
        assert note.path == "test.docx"

    def test_import_error_when_docx_missing(self):
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        with patch.object(builtins, "__import__", fake_import):
            from src.ingestion.docx_parser import parse_docx

            with pytest.raises(ImportError):
                parse_docx(b"fake docx", "test.docx")


# ---------------------------------------------------------------------------
# Audio parser
# ---------------------------------------------------------------------------


class TestAudioParser:
    def test_parse_audio_returns_fallback_note(self):
        """parse_audio no longer uses OpenAI Whisper — returns fallback note."""
        from src.ingestion.audio_parser import parse_audio

        note = parse_audio(b"fake audio", "test.mp3")
        assert note.title == "Test"
        assert "Transcription unavailable" in note.content


# ---------------------------------------------------------------------------
# Image parser
# ---------------------------------------------------------------------------


class TestImageParser:
    def test_import_error_when_pil_missing(self):
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name in ("PIL", "pytesseract"):
                raise ImportError(f"No module named '{name}'")
            return real_import(name, *args, **kwargs)

        with patch.object(builtins, "__import__", fake_import):
            from src.ingestion.image_parser import parse_image

            with pytest.raises(ImportError):
                parse_image(b"fake image", "test.png")


# ---------------------------------------------------------------------------
# Web parser
# ---------------------------------------------------------------------------


class TestWebParser:
    def test_parse_web_clip_success(self):
        from src.ingestion.web_parser import parse_web_clip

        html = (
            "<html><head><title>Article Title</title></head>"
            "<body><h1>Article Title</h1>"
            "<p>First paragraph.</p>"
            "<p>Second paragraph.</p>"
            "</body></html>"
        )
        with patch("requests.get") as mock_get:
            mock_get.return_value = MagicMock(
                status_code=200, text=html, raise_for_status=lambda: None
            )
            note = parse_web_clip("https://example.com/article")
        assert note.title == "Article Title"
        assert "First paragraph." in note.content
        assert note.path == "https://example.com/article"

    def test_import_error_when_requests_missing(self):
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "requests":
                raise ImportError("No module named 'requests'")
            return real_import(name, *args, **kwargs)

        with patch.object(builtins, "__import__", fake_import):
            from src.ingestion.web_parser import parse_web_clip

            with pytest.raises(ImportError):
                parse_web_clip("https://example.com")

    def test_runtime_error_on_bad_url(self):
        from src.ingestion.web_parser import parse_web_clip

        with patch("requests.get") as mock_get:
            mock_get.side_effect = Exception("Connection timeout")
            with pytest.raises(RuntimeError):
                parse_web_clip("https://example.com/article")


# ---------------------------------------------------------------------------
# Email parser
# ---------------------------------------------------------------------------


class TestEmailParser:
    def test_parse_email_simple(self):
        from src.ingestion.email_parser import parse_email

        raw = (
            b"Subject: Hello World\r\n"
            b"Date: Mon, 01 Jan 2024 12:00:00 +0000\r\n"
            b"\r\n"
            b"This is the body."
        )
        note = parse_email(raw, "test.eml")
        assert note.title == "Hello World"
        assert "This is the body." in note.content
        assert note.created.year == 2024

    def test_parse_email_multipart(self):
        from src.ingestion.email_parser import parse_email

        raw = (
            b"Subject: Multi\r\n"
            b"MIME-Version: 1.0\r\n"
            b"Content-Type: multipart/alternative; boundary=bound\r\n"
            b"\r\n"
            b"--bound\r\n"
            b"Content-Type: text/plain\r\n"
            b"\r\n"
            b"Plain text\r\n"
            b"--bound\r\n"
            b"Content-Type: text/html\r\n"
            b"\r\n"
            b"<html><body>HTML</body></html>\r\n"
            b"--bound--\r\n"
        )
        note = parse_email(raw, "multi.eml")
        assert note.title == "Multi"
        assert "Plain text" in note.content
        assert "HTML" in note.content

    def test_parse_email_bad_input_graceful(self):
        from src.ingestion.email_parser import parse_email

        note = parse_email(b"not an email", "bad.eml")
        assert isinstance(note.title, str)
        assert note.content == "not an email"
